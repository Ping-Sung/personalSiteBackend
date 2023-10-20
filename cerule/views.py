from django.shortcuts import render
from rest_framework.decorators import api_view

from django.http import JsonResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import time


@api_view(['POST'])
def newDoc(request, *args, **kwargs):
    data = request.data

    name = data.get('name', '無')
    sex = data.get('sex')
    email = data.get('email', '無')
    line = data.get('line', '無')
    phone = data.get('phone', '無')
    age = data.get('age')
    address = data.get('address', '無')
    note = data.get('note')
    if note == '':
        note = '\n\n\n\n'

    document = Document()
    heading = document.add_heading(level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run('Cerule 潛在客戶資料')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    table = document.add_table(rows=1, cols=3)
    first = table.rows[0].cells
    first[0].text = '名字：' + name
    first[1].text = '年齡：' + age
    first[2].text = '性別：' + sex

    second = table.add_row().cells
    second[0].text = 'Email：' + email
    second[2].text = 'Line ID：' + line
    a = second[0]
    b = second[1]
    A = a.merge(b)

    third = table.add_row().cells
    third[0].text = '地址：' + address
    third[2].text = '電話：' + phone
    a = third[0]
    b = third[1]
    A = a.merge(b)

    table = document.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '備註：' + note

    progress = document.add_paragraph()
    font_styles = document.styles
    font_charstyle = font_styles.add_style(
        'CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    # font_object.name = '楷體-繁'
    font_object.size = Pt(16)
    progress = document.add_paragraph()
    progress.add_run('進度追蹤', style='CommentsStyle').bold = True

    table = document.add_table(rows=1, cols=4)
    first = table.rows[0].cells

    paragraph = first[0].paragraphs[0]
    run = paragraph.add_run("□ 已撥打電話")
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    paragraph = first[1].paragraphs[0]
    run = paragraph.add_run("□ 已購買產品")
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    paragraph = first[2].paragraphs[0]
    run = paragraph.add_run("□ 已參加會議")
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    paragraph = first[3].paragraphs[0]
    run = paragraph.add_run("□ 拒絕聯絡")
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    notice = document.add_paragraph()
    notice.add_run('注意事項', style='CommentsStyle').bold = True
    notice = document.add_paragraph(
        style='List Number'
    )
    run = notice.add_run('本表格僅供美商施樂恩及其合作夥伴使用，請勿外流。')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    notice = document.add_paragraph(
        style='List Number'
    )
    run = notice.add_run('若有問題請洽 鄭博燦：0912345678 林秀玲：0912345567。')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    notice = document.add_paragraph(
        style='List Number'
    )
    run = notice.add_run('不知道還要寫什麼免責聲明...')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    document.add_page_break()
    fileDir = 'static/' + name + '_' + str(time.time()) + '.docx'
    document.save(fileDir)

    #   delete the file after 10 minutes

    return JsonResponse({'name': name}, status=200)
